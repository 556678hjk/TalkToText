import os
import sys
import tkinter as tk
from tkinter import filedialog, messagebox, ttk
import threading
import tempfile

import whisper
from inaSpeechSegmenter import Segmenter
from moviepy.editor import VideoFileClip
from docx import Document

# PyInstaller 路徑修正
base_path = os.path.dirname(os.path.abspath(sys.executable if getattr(sys, 'frozen', False) else __file__))

# 初始化 inaSpeechSegmenter
segmenter = Segmenter()

model_mapping = {
    "tiny(最快)": "tiny",
    "medium(次之)": "medium",
    # "large-v3(最慢)": "large-v3"
}

# Helper 函式：將秒數轉換成 hh:mm:ss 格式
def seconds_to_hhmmss(seconds):
    hours = int(seconds // 3600)
    minutes = int((seconds % 3600) // 60)
    secs = int(seconds % 60)
    return f"{hours:02d}:{minutes:02d}:{secs:02d}"

def process_audio_file(file_path, update_progress, model_name):
    temp_audio_file = None
    ext = os.path.splitext(file_path)[1].lower()
    try:
        update_progress(0)
        if ext == ".mp4":
            video = VideoFileClip(file_path)
            temp_audio_file = tempfile.mktemp(suffix=".wav")
            video.audio.write_audiofile(temp_audio_file, codec="pcm_s16le", logger=None)
            file_to_process = temp_audio_file
        else:
            file_to_process = file_path

        update_progress(10)

        whisper_model_path = os.path.join(base_path, 'whisper', 'assets')
        asr_model = whisper.load_model(model_name, download_root=whisper_model_path)

        result = asr_model.transcribe(file_to_process, language="zh") # 語言設定為中文
        whisper_segments = result.get("segments", [])

        update_progress(60)
        segmentation = segmenter(file_to_process)
        update_progress(100)

        return segmentation, whisper_segments

    finally:
        if temp_audio_file and os.path.exists(temp_audio_file):
            os.remove(temp_audio_file)

def start_processing():
    file_path = file_path_var.get()
    model_name = model_var.get()

    if not file_path:
        messagebox.showwarning("警告", "請先選擇音訊/視訊檔案")
        return

    output_text.delete(1.0, tk.END)
    progress_bar.config(value=0)
    output_text.insert(tk.END, "正在處理，請稍候...\n")

    def run():
        def update_progress(val):
            root.after(0, lambda: progress_bar.config(value=val))

        try:
            segmentation, whisper_segments = process_audio_file(file_path, update_progress, model_name)
            global speaker_results
            speaker_results = []
            for seg in segmentation:
                label, seg_start, seg_end = seg
                segment_text = ""
                for ws in whisper_segments:
                    ws_start = ws["start"]
                    ws_end = ws["end"]
                    if ws_end >= seg_start and ws_start <= seg_end:
                        segment_text += ws["text"].strip() + " "
                start_str = seconds_to_hhmmss(seg_start)
                end_str = seconds_to_hhmmss(seg_end)
                speaker_results.append(
                    f"{label} ({start_str} - {end_str}): {segment_text.strip()}"
                )
            root.after(0, lambda: output_text.delete(1.0, tk.END))
            root.after(0, lambda: output_text.insert(tk.END, "【講者對應轉錄結果】\n"))
            for line in speaker_results:
                root.after(0, lambda l=line: output_text.insert(tk.END, l + "\n"))
            root.after(0, lambda: export_button.config(state=tk.NORMAL))
        except Exception as e:
            root.after(0, lambda: output_text.insert(tk.END, f"處理時發生錯誤: {str(e)}\n"))

    threading.Thread(target=run).start()

def export_to_word():
    if not speaker_results:
        messagebox.showwarning("警告", "沒有可匯出的內容")
        return
    file_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word 文件", "*.docx")])
    if not file_path:
        return
    doc = Document()
    doc.add_heading("講者對應轉錄結果", level=1)
    for line in speaker_results:
        doc.add_paragraph(line)
    doc.save(file_path)
    messagebox.showinfo("成功", "轉錄結果已成功匯出為 Word 文件！")

def reset():
    file_path_var.set("")
    file_label.config(text="未選擇檔案")
    output_text.delete(1.0, tk.END)
    progress_bar.config(value=0)
    export_button.config(state=tk.DISABLED)

def select_file():
    file_path = filedialog.askopenfilename(filetypes=[("音訊/視訊檔案", "*.wav *.mp3 *.m4a *.mp4")])
    if file_path:
        file_path_var.set(file_path)
        file_label.config(text=file_path)

def on_select(event):
    """當選擇 Combobox 時，更新 model_var 為實際名稱"""
    selected_display = model_dropdown.get()  # 取得選中的顯示名稱
    model_var.set(model_mapping[selected_display])  # 設定為實際模型名稱
    print(f"選擇的模型: {model_var.get()}")  # 測試輸出

root = tk.Tk()
root.title("語音轉文字與講者對應")
root.geometry("600x600")

file_path_var = tk.StringVar()

select_button = tk.Button(root, text="選擇音訊/視訊檔案", command=select_file)
select_button.pack(pady=10)

file_label = tk.Label(root, text="未選擇檔案")
file_label.pack()

model_var = tk.StringVar(value="tiny")
display_values = list(model_mapping.keys())
model_dropdown = ttk.Combobox(root, values=display_values, state="readonly")
model_dropdown.pack(pady=10)
model_dropdown.config(width=15)
model_dropdown.current(0)
model_var.set(model_mapping[display_values[0]])
model_dropdown.bind("<<ComboboxSelected>>", on_select)

progress_bar = ttk.Progressbar(root, orient="horizontal", length=400, mode="determinate", maximum=100)
progress_bar.pack(pady=10)

process_button = tk.Button(root, text="開始處理", command=start_processing)
process_button.pack(pady=10)

output_text = tk.Text(root, wrap=tk.WORD)
output_text.pack(expand=True, fill=tk.BOTH, padx=10, pady=10)

button_frame = tk.Frame(root)
button_frame.pack(pady=10)

export_button = tk.Button(button_frame, text="匯出 Word", command=export_to_word, state=tk.DISABLED)
export_button.pack(side=tk.LEFT, padx=5)

reset_button = tk.Button(button_frame, text="重設", command=reset)
reset_button.pack(side=tk.LEFT, padx=5)

root.mainloop()
